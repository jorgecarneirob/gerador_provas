#!/usr/bin/env python
# coding: utf-8

# In[19]:


import random
from docx import Document

# Variáveis globais
NUM_QUESTOES = 10
#ALTERNATIVAS = ["A", "B", "C", "D", "E"]
TIPOS_PROVA = [1, 2, 3, 4]

def obter_gabarito_original():
    gabarito_base = {}

    for q in range(1, NUM_QUESTOES + 1):
        print(f"\n=== QUESTÃO {q} ===")
        enunciado = input(f"Enunciado da questão {q}: ").strip()

        textos_alt = {}
        pesos_alt = {}
        for alt in ALTERNATIVAS:
            texto = input(f"Texto da alternativa {alt}: ").strip()
            while True:
                try:
                    peso = float(input(f"Peso da alternativa {alt}: ").strip())
                    break
                except ValueError:
                    print("Valor inválido. Digite um número (use ponto para decimais).")

            textos_alt[alt] = texto
            pesos_alt[alt] = peso

        correta = ""
        while True:
            correta = input(f"Alternativa correta (A-E): ").strip().upper()
            if correta in ALTERNATIVAS:
                break
            print("Alternativa inválida. Digite A, B, C, D ou E.")

        while True:
            try:
                peso_q = float(input(f"Peso da questão {q}: ").strip())
                break
            except ValueError:
                print("Valor inválido. Digite um número, use ponto para decimais.")


        gabarito_base[f"Q{q}"] = {
            "enunciado": enunciado,
            "alternativas": textos_alt,
            "pesos_alt": pesos_alt,
            "correta": correta,
            "peso_questao": peso_q
        }

    return gabarito_base

def gerar_tipos_prova(gabarito_base, tipos=TIPOS_PROVA):
    provas = {}

    for tipo in tipos:
        prova = {}
        for q, dados in gabarito_base.items():
            enunciado = dados["enunciado"]
            textos_alt = dados["alternativas"]
            pesos_alt = dados["pesos_alt"]
            correta = dados["correta"]
            peso_q = dados["peso_questao"]

           # Embaralhar lista de letras originais
            alt_embaralhadas = ALTERNATIVAS.copy()
            random.shuffle(alt_embaralhadas)

            # criar novas letras A,B,C,D,E
            novas_letras = ALTERNATIVAS.copy()

            # Montar novo dicionário com letras A-E
            textos_embaralhados = {}
            pesos_embaralhados = {}

            # Descobrir qual letra nova corresponde à correta
            nova_correta = None

            for i, alt_original in enumerate(alt_embaralhadas):
                letra_nova = novas_letras[i]
                textos_embaralhados[letra_nova] = textos_alt[alt_original]
                pesos_embaralhados[letra_nova] = pesos_alt[alt_original]

                if alt_original == correta:
                    nova_correta = letra_nova

            prova[q] = {
                "enunciado": enunciado,
                "alternativas": textos_embaralhados,
                "pesos_alt": pesos_embaralhados,
                "correta": nova_correta,
                "peso_questao": peso_q
            }


        provas[str(tipo)] = prova

    return provas

def salvar_gabarito_txt(provas, filename):
    linhas = []

    for tipo, prova in provas.items():
        for q, dados in sorted(prova.items(), key=lambda x: int(x[0][1:])):
            pesos_alt_txt = ",".join(
                [f"{alt}:{dados['pesos_alt'][alt]}" for alt in ALTERNATIVAS]
            )

            linha = f"{tipo}|{q}|{dados['peso_questao']}|{pesos_alt_txt}|{dados['correta']}"
            linhas.append(linha)

    with open(arquivo, "w", encoding="utf-8") as f:
        f.write("\n".join(linhas))

    print(f"\nArquivo '{filename}' gerado com sucesso!")

def gerar_word_unificado(provas, filename):
    doc = Document()

    for tipo, prova in provas.items():
        # Cabeçalho
        
        doc.add_paragraph("Nome: ____________________________________________N° _________")
        doc.add_paragraph("Sala: ________disciplina: _________________ Data:___/____/____")
        doc.add_paragraph(f"Prova Tipo {tipo}")
        doc.add_paragraph("")  # Espaço extra

        # Questões
        for q, dados in sorted(prova.items(), key=lambda x: int(x[0][1:])):
            doc.add_paragraph(f"{q}. {dados['enunciado']}")
            for alt in ALTERNATIVAS:
                texto_alt = dados['alternativas'][alt]
                doc.add_paragraph(f"   ({alt}) {texto_alt}")
            doc.add_paragraph("")  # Espaço entre questões

        # Separar provas por quebra de página
        doc.add_page_break()

    doc.save(arquivo)
    print(f"\nArquivo Word '{arquivo}' gerado com sucesso!")

# EXECUÇÃO
if __name__ == "__main__":
    gabarito_base = obter_gabarito_original()
    provas = gerar_tipos_prova(gabarito_base)
    salvar_gabarito_txt(provas)
    gerar_word_unificado(provas)


# In[41]:


import random
from docx import Document
from docx.enum.text import WD_BREAK
import os # Importar para manipulação de diretórios

ALTERNATIVAS = ["A", "B", "C", "D", "E"]

def ler_entrada_txt(filepath, alternativas_list): # Alterado para aceitar filepath
    """
    Lê o arquivo entrada.txt e gera o gabarito-base
    """
    gabarito_base = {}
    with open(filepath, "r", encoding="utf-8") as f:
        lines = [l.rstrip("\n") for l in f if l.strip() != ""]

    i = 0
    while i < len(lines):
        # Q1|1.0
        cab = lines[i]
        q_id, peso_q = cab.split("|")
        peso_q = float(peso_q)
        i += 1

        enunciado = lines[i]
        i += 1

        alternativas = {}
        pesos_alt = {}
        # Usar alternativas_list para iterar
        for _ in range(len(alternativas_list)):
            letra, texto, peso = lines[i].split("|")
            alternativas[letra] = texto
            pesos_alt[letra] = float(peso)
            i += 1

        # A alternativa correta ainda é determinada pelo maior peso
        correta = max(pesos_alt, key=pesos_alt.get)
        gabarito_base[q_id] = {
            "enunciado": enunciado,
            "alternativas": alternativas,
            "pesos_alt": pesos_alt,
            "correta": correta,
            "peso_questao": peso_q
        }
    return gabarito_base

def gerar_provas(gabarito_base, num_tipos, alternativas_list): # Adicionado alternativas_list e num_tipos
    """
    Gera várias provas embaralhando as alternativas
    """
    todas_provas = {}
    for tipo in range(1, num_tipos + 1): # Usar num_tipos
        prova = {}
        for q_id, dados in gabarito_base.items():
            enunciado = dados["enunciado"]
            alternativas = dados["alternativas"]
            pesos_alt = dados["pesos_alt"]

            letras = list(alternativas.keys())
            random.shuffle(letras)

            alt_embaralhadas = {}
            pesos_embaralhados = {}
            # Usar alternativas_list para zip
            for letra_nova, letra_old in zip(alternativas_list, letras):
                alt_embaralhadas[letra_nova] = alternativas[letra_old]
                pesos_embaralhados[letra_nova] = pesos_alt[letra_old]

            # nova correta:
            nova_correta = max(pesos_embaralhados, key=pesos_embaralhados.get)

            prova[q_id] = {
                "enunciado": enunciado,
                "alternativas": alt_embaralhadas,
                "pesos_alt": pesos_embaralhados,
                "correta": nova_correta,
                "peso_questao": dados["peso_questao"]
            }
        todas_provas[str(tipo)] = prova
    return todas_provas

def salvar_gabarito_arquivo(provas, filename, alternativas_list): # Adicionado alternativas_list
    """
    Salva todas as provas em arquivo gabarito.txt
    """
    with open(filename, "w", encoding="utf-8") as f:
        for tipo, prova in provas.items():
            for q_id, dados in prova.items():
                peso_q = dados["peso_questao"]
                # Usar os pesos das alternativas que já estão no dicionário de dados
                alt_pesos = ",".join(f"{alt}:{dados['pesos_alt'][alt]}" for alt in alternativas_list)
                correta = dados["correta"]
                f.write(f"{tipo}|{q_id}|{peso_q}|{alt_pesos}|{correta}\n")

    print(f"\nArquivo '{filename}' gerado com sucesso!")

def gerar_word_unificado(provas, filename, alternativas_list): # Adicionado alternativas_list
    """
    Gera um arquivo Word único com as provas
    """
    doc = Document()

    for tipo, prova in provas.items():
        doc.add_paragraph("Nome: __________________________________________________N° _________")
        doc.add_paragraph("Sala: ________disciplina: _______________________ Data:___/____/____")
        doc.add_paragraph(f"Prova Tipo {tipo}")
        doc.add_paragraph("")  # Espaço extra

        for q_id, dados in sorted(prova.items(), key=lambda x: int(x[0][1:])):
            doc.add_paragraph(f"{q_id} - {dados['enunciado']}", style='Normal')
            # Usar alternativas_list para iterar
            for alt in alternativas_list:
                texto = dados["alternativas"][alt]
                doc.add_paragraph(f"({alt}) {texto}", style='Normal')
            doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    doc.save(filename)
    print(f"Arquivo Word salvo: {filename}")

# Removido o bloco if __name__ == "__main__": pois o Flask chamará as funções diretamente.
# As funções agora aceitam os caminhos dos arquivos como argumentos.